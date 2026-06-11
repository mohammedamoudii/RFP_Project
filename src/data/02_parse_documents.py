from pathlib import Path
import json
import pandas as pd
from pypdf import PdfReader
from docx import Document
from tqdm import tqdm


MANIFEST_PATH = Path("data/processed/file_manifest.csv")
OUTPUT_PATH = Path("data/processed/parsed_elements.jsonl")
ERRORS_PATH = Path("data/processed/parsing_errors.csv")


def write_jsonl(rows, output_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with output_path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def extract_pdf(row: dict):
    file_path = Path(row["absolute_path"])
    reader = PdfReader(str(file_path))

    elements = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""

        if text.strip():
            elements.append(
                {
                    "doc_id": row["file_id"],
                    "project_id": row["project_id"],
                    "project_name": row["project_name"],
                    "source_file": row["source_file"],
                    "file_type": row["file_type"],
                    "relative_path": row["relative_path"],
                    "absolute_path": row["absolute_path"],
                    "page_number": page_index,
                    "element_index": page_index,
                    "element_type": "pdf_page",
                    "section": None,
                    "content": text,
                }
            )

    return elements


def extract_docx(row: dict):
    file_path = Path(row["absolute_path"])
    doc = Document(str(file_path))

    elements = []
    element_index = 0

    # Extract normal paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        element_index += 1

        style_name = paragraph.style.name if paragraph.style else ""

        if "heading" in style_name.lower():
            element_type = "heading"
        else:
            element_type = "paragraph"

        elements.append(
            {
                "doc_id": row["file_id"],
                "project_id": row["project_id"],
                "project_name": row["project_name"],
                "source_file": row["source_file"],
                "file_type": row["file_type"],
                "relative_path": row["relative_path"],
                "absolute_path": row["absolute_path"],
                "page_number": None,
                "element_index": element_index,
                "element_type": element_type,
                "section": text if element_type == "heading" else None,
                "content": text,
            }
        )

    # Extract tables
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, table_row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in table_row.cells]
            cells = [cell for cell in cells if cell]

            if not cells:
                continue

            element_index += 1
            text = " | ".join(cells)

            elements.append(
                {
                    "doc_id": row["file_id"],
                    "project_id": row["project_id"],
                    "project_name": row["project_name"],
                    "source_file": row["source_file"],
                    "file_type": row["file_type"],
                    "relative_path": row["relative_path"],
                    "absolute_path": row["absolute_path"],
                    "page_number": None,
                    "element_index": element_index,
                    "element_type": f"docx_table_{table_index}_row_{row_index}",
                    "section": None,
                    "content": text,
                }
            )

    return elements


def parse_file(row: dict):
    file_type = str(row["file_type"]).lower()

    if file_type == "pdf":
        return extract_pdf(row)

    if file_type == "docx":
        return extract_docx(row)

    return []


def main():
    if not MANIFEST_PATH.exists():
        raise FileNotFoundError(
            f"Manifest file not found: {MANIFEST_PATH}. Run 01_create_manifest.py first."
        )

    manifest = pd.read_csv(MANIFEST_PATH)

    all_elements = []
    errors = []

    for _, row in tqdm(manifest.iterrows(), total=len(manifest), desc="Parsing files"):
        row_dict = row.to_dict()

        try:
            elements = parse_file(row_dict)
            all_elements.extend(elements)

        except Exception as e:
            errors.append(
                {
                    "file_id": row_dict.get("file_id"),
                    "project_name": row_dict.get("project_name"),
                    "source_file": row_dict.get("source_file"),
                    "file_type": row_dict.get("file_type"),
                    "relative_path": row_dict.get("relative_path"),
                    "absolute_path": row_dict.get("absolute_path"),
                    "error_message": str(e),
                }
            )

    write_jsonl(all_elements, OUTPUT_PATH)

    errors_df = pd.DataFrame(errors)
    errors_df.to_csv(ERRORS_PATH, index=False, encoding="utf-8-sig")

    print(f"Parsed elements saved to: {OUTPUT_PATH}")
    print(f"Parsing errors saved to: {ERRORS_PATH}")
    print(f"Files in manifest: {len(manifest)}")
    print(f"Parsed text elements: {len(all_elements)}")
    print(f"Files with errors: {len(errors)}")

    if all_elements:
        print("\nSample parsed element:")
        sample = all_elements[0]
        print("Project:", sample["project_name"])
        print("Source:", sample["source_file"])
        print("Type:", sample["element_type"])
        print("Content preview:", sample["content"][:500])


if __name__ == "__main__":
    main()